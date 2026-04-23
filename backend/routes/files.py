import uuid
import re
import io
from datetime import datetime
from typing import Optional

from fastapi import APIRouter, Depends, UploadFile, File, Form
from database import transactions_collection
from routes.auth import get_current_user
from ml.classifier import classify_transaction

router = APIRouter()

# ── Sensitive data patterns ────────────────────────────────────────
PAN_RE        = re.compile(r'\b[A-Z]{5}[0-9]{4}[A-Z]\b')
AADHAAR_RE    = re.compile(r'\b[2-9]\d{3}[\s\-]?\d{4}[\s\-]?\d{4}\b')
CC_RE         = re.compile(r'\b(?:\d{4}[\s\-]?){3}\d{4}\b')
PHONE_RE      = re.compile(r'\b[6-9]\d{9}\b')
EMAIL_RE      = re.compile(r'\b[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,}\b')
IFSC_RE       = re.compile(r'\b[A-Z]{4}0[A-Z0-9]{6}\b')
PASSPORT_RE   = re.compile(r'\b[A-Z][1-9][0-9]{7}\b')

MALWARE_SIGS = [
    "powershell -enc", "powershell -e ", "cmd.exe /c", "regsvr32 /s",
    "mshta.exe", "wscript.exe", "cscript.exe", "certutil -decode",
    "bitsadmin /transfer", "shell_exec(", "eval(base64_decode",
    "CreateObject(\"WScript", "net user /add", "net localgroup administrators",
    "HKEY_LOCAL_MACHINE\\SAM", "mimikatz", "invoke-expression",
]

SENSITIVE_KEYWORDS = [
    "password", "passwd", "secret key", "private key", "api key", "bearer token",
    "confidential", "do not share", "restricted", "classified",
    "salary slip", "payroll", "bank statement", "account number", "ifsc code",
    "credit limit", "cvv", "pin number", "otp", "passphrase",
]


def extract_text(content_bytes: bytes, filename: str) -> str:
    """Extract text from PDF, DOCX, or plain text files."""
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else "txt"
    text = ""

    if ext == "pdf":
        try:
            import PyPDF2
            reader = PyPDF2.PdfReader(io.BytesIO(content_bytes))
            for page in reader.pages:
                try:
                    text += (page.extract_text() or "") + "\n"
                except Exception:
                    pass
        except Exception:
            text = content_bytes.decode("utf-8", errors="ignore")

    elif ext in ("docx", "doc"):
        try:
            import docx
            doc = docx.Document(io.BytesIO(content_bytes))
            for para in doc.paragraphs:
                text += para.text + "\n"
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                text += "\n"
        except Exception:
            text = content_bytes.decode("utf-8", errors="ignore")

    else:
        text = content_bytes.decode("utf-8", errors="ignore")

    return text


def analyze_content(text: str, filename: str) -> dict:
    """Full content analysis — PII, malware, sensitive keywords."""
    findings = {
        "sensitive_data": [],
        "malware_patterns": [],
        "sensitive_keywords_found": [],
        "pii_detected": False,
        "extracted_text_length": len(text),
        "content_risk_score": 0,
        "risk_reasons": [],
    }

    if not text.strip():
        return findings

    # PAN Card
    pans = list(set(PAN_RE.findall(text)))
    if pans:
        findings["sensitive_data"].append({
            "type": "PAN Card", "count": len(pans),
            "severity": "high", "examples": pans[:2],
            "description": f"Found {len(pans)} PAN card number(s) — may violate DPDP Act 2023"
        })
        findings["pii_detected"] = True
        findings["risk_reasons"].append(f"Contains {len(pans)} PAN card number(s)")

    # Aadhaar
    aadhaars = AADHAAR_RE.findall(text)
    if aadhaars:
        findings["sensitive_data"].append({
            "type": "Aadhaar Number", "count": len(aadhaars),
            "severity": "high", "examples": [],
            "description": f"Found {len(aadhaars)} Aadhaar-like number(s) — sensitive biometric ID"
        })
        findings["pii_detected"] = True
        findings["risk_reasons"].append(f"Contains {len(aadhaars)} Aadhaar number(s)")

    # Credit/Debit Cards
    cards = CC_RE.findall(text)
    if cards:
        findings["sensitive_data"].append({
            "type": "Credit/Debit Card", "count": len(cards),
            "severity": "critical", "examples": [],
            "description": f"Found {len(cards)} card number pattern(s) — PCI-DSS violation risk"
        })
        findings["pii_detected"] = True
        findings["risk_reasons"].append(f"Contains {len(cards)} card number(s) — critical PCI-DSS risk")

    # IFSC Codes
    ifscs = list(set(IFSC_RE.findall(text)))
    if ifscs:
        findings["sensitive_data"].append({
            "type": "Bank IFSC Code", "count": len(ifscs),
            "severity": "medium", "examples": ifscs[:2],
            "description": f"Found {len(ifscs)} bank IFSC code(s)"
        })
        findings["risk_reasons"].append("Contains bank IFSC code(s)")

    # Phone numbers (bulk = suspicious)
    phones = PHONE_RE.findall(text)
    if len(phones) > 3:
        findings["sensitive_data"].append({
            "type": "Phone Numbers (Bulk)", "count": len(phones),
            "severity": "medium", "examples": [],
            "description": f"Found {len(phones)} Indian phone numbers — possible customer data leak"
        })
        findings["pii_detected"] = True

    # Email addresses (bulk)
    emails = EMAIL_RE.findall(text)
    if len(emails) > 5:
        findings["sensitive_data"].append({
            "type": "Email Addresses (Bulk)", "count": len(emails),
            "severity": "medium", "examples": [],
            "description": f"Found {len(emails)} email addresses — possible contact list exposure"
        })

    # Sensitive keywords
    txt_lower = text.lower()
    found_kw = [kw for kw in SENSITIVE_KEYWORDS if kw in txt_lower]
    if found_kw:
        findings["sensitive_keywords_found"] = found_kw[:10]
        findings["risk_reasons"].append(f"Contains sensitive keywords: {', '.join(found_kw[:3])}")

    # Malware signatures
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    if ext in ("ps1", "bat", "cmd", "vbs", "js", "sh", "txt", "py"):
        found_mal = [sig for sig in MALWARE_SIGS if sig.lower() in txt_lower]
        if found_mal:
            findings["malware_patterns"] = found_mal[:5]
            findings["risk_reasons"].append(f"Malware patterns detected: {', '.join(found_mal[:2])}")

    # Calculate content risk score
    score = 0
    for item in findings["sensitive_data"]:
        if item["severity"] == "critical": score += 45
        elif item["severity"] == "high":   score += 30
        elif item["severity"] == "medium": score += 12
    score += len(findings["sensitive_keywords_found"]) * 4
    score += len(findings["malware_patterns"]) * 25
    findings["content_risk_score"] = min(100, score)

    return findings


@router.post("/send")
async def send_file(
    recipient_email: str = Form(...),
    subject: str = Form(...),
    body: str = Form(...),
    file: Optional[UploadFile] = File(None),
    current_user=Depends(get_current_user),
):
    filename = ""
    file_size = 0
    content_analysis = None

    if file and file.filename:
        filename = file.filename
        content_bytes = await file.read()
        file_size = len(content_bytes)

        # Deep content analysis for files under 10MB
        if file_size < 10 * 1024 * 1024:
            extracted_text = extract_text(content_bytes, filename)
            content_analysis = analyze_content(extracted_text, filename)

    # ML classification on email metadata
    result = classify_transaction(subject, body, filename)

    # Blend content risk with ML risk
    final_risk = result["risk_score"]
    if content_analysis and content_analysis["content_risk_score"] > 0:
        content_boost = content_analysis["content_risk_score"] * 0.45
        final_risk = min(100, round(final_risk + content_boost, 1))
        result["risk_score"] = final_risk
        if final_risk >= 70:
            result["severity"] = "high"
            result["classification"] = "suspicious"
        elif final_risk >= 35:
            result["severity"] = "medium"

    # Build explanation
    risk_reasons = result.get("suspicious_keywords", [])
    if content_analysis:
        risk_reasons = content_analysis.get("risk_reasons", []) + risk_reasons

    transaction = {
        "_id": str(uuid.uuid4()),
        "sender_id": current_user["_id"],
        "sender_name": current_user["name"],
        "sender_email": current_user.get("email") or current_user.get("phone") or "",
        "recipient_email": recipient_email,
        "subject": subject,
        "body": body[:500],
        "filename": filename,
        "file_size": file_size,
        "classification": result["classification"],
        "risk_score": result["risk_score"],
        "severity": result["severity"],
        "suspicious_keywords": result["suspicious_keywords"],
        "content_analysis": content_analysis,
        "risk_reasons": risk_reasons[:8],
        "timestamp": datetime.utcnow(),
        "type": "file_transaction",
    }

    await transactions_collection.insert_one(transaction)

    return {
        "success": True,
        "transaction_id": transaction["_id"],
        "classification": result["classification"],
        "risk_score": result["risk_score"],
        "severity": result["severity"],
        "suspicious_keywords": result["suspicious_keywords"],
        "content_analysis": content_analysis,
        "risk_reasons": risk_reasons[:8],
    }


@router.get("/my-history")
async def my_history(current_user=Depends(get_current_user)):
    cursor = transactions_collection.find(
        {"sender_id": current_user["_id"], "type": "file_transaction"},
        sort=[("timestamp", -1)],
    ).limit(50)
    results = []
    async for t in cursor:
        t["_id"] = str(t["_id"])
        t["timestamp"] = t["timestamp"].isoformat() if t.get("timestamp") else ""
        results.append(t)
    return results